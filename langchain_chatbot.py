import streamlit as st
import os
import pandas as pd
from docx import Document as DocxDocument
#pip install python-docx
from langchain.schema import Document as LangchainDocument
from langchain.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
import requests

# Initialize chat history
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Set page config
st.set_page_config(
    page_title="Denmark Competitors Q&A",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .stTextInput > div > div > input {
        background-color: #f0f2f6;
    }
    .stButton > button {
        width: 100%;
        background-color: #FF4B4B;
        color: white;
        border: none;
        padding: 10px 24px;
        border-radius: 5px;
        font-weight: bold;
    }
    .stButton > button:hover {
        background-color: #FF6B6B;
    }
    </style>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.title("🚀 About")
    st.markdown("""
    This application allows you to identify potential competitors among Consistent high-growth firms in different sectors of Denmark.
    The data is analyzed using advanced AI technology to provide accurate and relevant answers.
    """)
    
    st.markdown("---")
    st.markdown("### 💡 Tips")
    st.markdown("""
    - Ask specific questions about competitors
    - Include relevant context in your questions
    - Try different phrasings if you don't get the desired answer
    """)

# Main content
st.title("🤖 Denmark Competitors Q&A Assistant")
st.markdown("Ask questions about potential competitors in different sectors and get AI-powered answers!")

# Toggleable Sidebar to use Web Search 
use_web = st.sidebar.checkbox("🌐 Use Web Search (Serper)", value=True)

# Load data and initialize components
@st.cache_resource
def load_data():
    # Load API key
    with open(r'D:\STUDIA\SGH_mag\SEMESTR_4_Erasmus\Data-Driven\groq_api_key.txt', 'r', encoding='utf-8') as file:
        api_key = str(file.readline())
    os.environ['GROQ_API_KEY'] = api_key

    # Load and clean CSV
    df = pd.read_csv('topicanalysis_denmark_few.csv')
    df = df.drop(columns=[col for col in df.columns if 'Unnamed' in col or df[col].isnull().all()])
    df.fillna('', inplace=True)

    # Combine relevant columns for semantic search
    df['combined_text'] = (
        df['Company name Latin alphabet'] + " - " +
        df['City Latin Alphabet'] + " " +
        df['Region in country'] + " " +
        df['Standardized legal form'] + " " +
        df['Final Description'] + " " +
        df['Topics']
    )

    # Convert to LangChain documents (to improve RAG relevance)
    documents = [
        LangchainDocument(
            page_content=row['combined_text'],
            metadata={
                "company": row['Company name Latin alphabet'],
                "founded_year": row['Founded Year'],
                "sector": row['NACE Rev. 2 main section'],
                "n_emp": row['Number of employees 2023'],
                "growth": row['Growth 2023']
            }
        )
        for _, row in df.iterrows()
    ]

    # Load embeddings
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"}
    )

    # Build FAISS vector store
    vectorstore = FAISS.from_documents(documents, embeddings)

    # Initialize LLM
    llm = ChatGroq(model="qwen-qwq-32b")

    # System prompt 
    custom_prompt = PromptTemplate.from_template("""
    You are an expert business analyst helping identify potential competitors among high-growth firms in Denmark.

    The user is expected to give a company name and description. Based on that:
    - Return the **top 3 most relevant competitors** from the provided context.
    - For each competitor, explain **why** it’s a strong match (based on sector, description, region, number of employees, growth, or business model).
    - Be concise, clear, and use bullet points for clarity.

    Context:
    {context}

    User Question:
    {question}
    """)

    # Return conversational chain
    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        verbose=True
    )

# Load the chain
placeholder = st.empty()
with placeholder.container():
    st.write("⚙️ Initializing AI backend...")
    chain = load_data()
placeholder.empty()  # Clear once loaded

# Display previous messages
if 'chat_history' in st.session_state:
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# Use Web Search with Serper API
def serper_search(query: str, api_key: str) -> list:
    url = "https://google.serper.dev/search"
    headers = {"X-API-KEY": api_key}
    payload = {"q": query, "num": 3} # limit the results

    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        results = response.json().get("organic", [])
        return [
            f"- {r['title']}: {r.get('snippet', '')} ({r['link']})"
            for r in results
        ]
    except Exception as e:
        return [f"Web search failed: {e}"]

# Generate DOCX
def generate_docx_report(text: str, file_path: str):
    doc = DocxDocument()
    doc.add_heading("Competitor Analysis Report", level=1)

    for line in text.strip().split('\n'):
        doc.add_paragraph(line)

    doc.save(file_path)

with open(r'D:\STUDIA\SGH_mag\SEMESTR_4_Erasmus\Data-Driven\serper_api_key.txt', 'r', encoding='utf-8') as file:
        SERPER_API_KEY = str(file.readline())
os.environ['SERPER_API_KEY'] = SERPER_API_KEY

# Chat input
if prompt := st.chat_input("Ask your question about competitors..."):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
    
    # Get AI response
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            web_context = ""
            if use_web:
                web_results = serper_search(prompt, SERPER_API_KEY)
                web_context = "\n\nWeb Search Results:\n" + "\n".join(web_results)
            full_query = f"{prompt}{web_context}" # add web context to the query as Langchain uses only FAISS
            result = chain.invoke({"question": full_query, "chat_history": []})
            response = result['answer']
            st.markdown(response)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            st.session_state.last_response = response  # Save for download

            if response:
                file_path = "competitor_analysis_report.docx"
                generate_docx_report(response, file_path)

                with open(file_path, "rb") as f:
                    st.download_button(
                    label="📝 Download Report as Word (.docx)",
                    data=f,
                    file_name="competitor_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

# Sidebar: clear chat history
with st.sidebar:
    if st.button("🗑️ Clear Chat History"):
        st.session_state.pop("chat_history", None)
        st.experimental_rerun()

# Footer
st.markdown("---")
st.markdown("""
    <div style='text-align: center'>
        <p>Built with ❤️ using Streamlit and LangChain</p>
        <p>Powered by Groq AI</p>
    </div>
""", unsafe_allow_html=True)
